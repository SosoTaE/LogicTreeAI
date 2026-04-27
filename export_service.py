"""
Export multi-agent discussion sessions as Word (.docx) or PDF.

Both formats render the same logical document:
  - Title (session title)
  - Metadata block (mode, status, rounds, participants, timestamps)
  - Problem Statement section
  - One section per turn: "[Turn N] model_name (role?)" + content, errors flagged
"""
import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_LEFT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak, Preformatted,
)
from reportlab.platypus.flowables import HRFlowable


def _safe_filename_stub(title):
    """Sanitize a session title into a safe filename stem."""
    cleaned = ''.join(
        c if (c.isalnum() or c in ('-', '_', ' ')) else '_'
        for c in (title or 'discussion')
    ).strip()
    cleaned = '_'.join(cleaned.split())
    return cleaned[:80] or 'discussion'


def export_filename(session, extension):
    stub = _safe_filename_stub(getattr(session, 'title', None))
    stamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
    return f'{stub}_{stamp}.{extension}'


def _iter_turns_in_order(session):
    """
    Return turns ordered by (turn_number, created_at) so parallel rounds
    keep a stable, readable order.
    """
    return sorted(
        session.turns,
        key=lambda t: (t.turn_number, t.created_at or datetime.min),
    )


def _metadata_lines(session):
    participants = session.participating_models or []
    lines = [
        ('Mode', (session.conversation_mode or '').capitalize()),
        ('Status', (session.status or '').capitalize()),
        ('Turn progress', f'{session.current_round} / {session.max_rounds}'),
        ('Participants', ', '.join(participants) if participants else '—'),
    ]
    if session.created_at:
        lines.append(('Started', session.created_at.isoformat(timespec='seconds')))
    if session.completed_at:
        lines.append(('Ended', session.completed_at.isoformat(timespec='seconds')))
    return lines


# ------------------------- Word (.docx) -------------------------

def session_to_docx(session):
    """Render the session to a .docx file and return the bytes."""
    doc = Document()

    # Overall page margins
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # Title
    title_para = doc.add_heading(session.title or 'Multi-Agent Discussion', level=0)
    for run in title_para.runs:
        run.font.size = Pt(22)

    # Metadata
    meta_table = doc.add_table(rows=0, cols=2)
    meta_table.autofit = True
    for label, value in _metadata_lines(session):
        row = meta_table.add_row().cells
        lbl = row[0].paragraphs[0].add_run(label)
        lbl.bold = True
        lbl.font.size = Pt(10)
        val_run = row[1].paragraphs[0].add_run(str(value))
        val_run.font.size = Pt(10)

    doc.add_paragraph()

    # Problem statement
    doc.add_heading('Problem Statement', level=1)
    problem_para = doc.add_paragraph(session.initial_problem or '')
    for run in problem_para.runs:
        run.font.size = Pt(11)

    # Discussion turns
    doc.add_heading('Discussion', level=1)
    turns = _iter_turns_in_order(session)
    if not turns:
        doc.add_paragraph('(No turns recorded.)')
    else:
        for turn in turns:
            name_lower = (turn.model_name or '').lower()
            is_user = name_lower == 'user'
            is_moderator = name_lower == 'moderator'

            if is_moderator:
                # Render moderator decisions as a slim italic line
                # rather than a full heading + body block — they're
                # meta-events, not part of the substantive discussion.
                chosen = (turn.model_role or '').lstrip('->').strip()
                summary = f'Moderator -> {chosen}' if chosen else 'Moderator'
                mod_para = doc.add_paragraph()
                mod_run = mod_para.add_run(summary)
                mod_run.bold = True
                mod_run.italic = True
                mod_run.font.size = Pt(10)
                mod_run.font.color.rgb = RGBColor(0x55, 0x55, 0x88)
                if turn.content and turn.content.strip():
                    reason = doc.add_paragraph()
                    reason_run = reason.add_run(turn.content.strip())
                    reason_run.italic = True
                    reason_run.font.size = Pt(9.5)
                    reason_run.font.color.rgb = RGBColor(0x66, 0x66, 0x77)
                continue

            label = 'Human (user)' if is_user else (turn.model_name or 'unknown')
            heading_bits = [f'Turn {turn.turn_number}', label]
            if turn.model_role and not is_user:
                heading_bits.append(f'role: {turn.model_role}')
            if turn.duration is not None and not is_user:
                heading_bits.append(f'{turn.duration:.2f}s')
            heading = ' · '.join(heading_bits)

            h = doc.add_heading(heading, level=2)
            for run in h.runs:
                run.font.size = Pt(13)

            if turn.error:
                err_para = doc.add_paragraph()
                err_run = err_para.add_run(f'[error] {turn.error}')
                err_run.font.color.rgb = RGBColor(0xB0, 0x00, 0x20)
                err_run.italic = True
                err_run.font.size = Pt(10)
            else:
                # Split on blank lines so paragraphs survive.
                text = turn.content or ''
                for block in text.split('\n\n'):
                    if not block.strip():
                        continue
                    p = doc.add_paragraph(block.strip())
                    for run in p.runs:
                        run.font.size = Pt(11)

    # Synthesis section (if available)
    if session.synthesis:
        doc.add_heading('Synthesis', level=1)
        if session.synthesis_model:
            model_para = doc.add_paragraph()
            model_run = model_para.add_run(f'Synthesized by: {session.synthesis_model}')
            model_run.italic = True
            model_run.font.size = Pt(10)

        # Split on blank lines for better formatting
        for block in session.synthesis.split('\n\n'):
            if not block.strip():
                continue
            p = doc.add_paragraph(block.strip())
            for run in p.runs:
                run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ------------------------- PDF -------------------------

def _pdf_styles():
    base = getSampleStyleSheet()
    styles = {
        'Title': ParagraphStyle(
            'DiscussTitle', parent=base['Title'],
            fontName='Helvetica-Bold', fontSize=20,
            spaceAfter=14, textColor=HexColor('#1b1f3a'),
            alignment=TA_LEFT,
        ),
        'H1': ParagraphStyle(
            'DiscussH1', parent=base['Heading1'],
            fontName='Helvetica-Bold', fontSize=14,
            spaceBefore=16, spaceAfter=8, textColor=HexColor('#1b1f3a'),
        ),
        'Turn': ParagraphStyle(
            'DiscussTurn', parent=base['Heading2'],
            fontName='Helvetica-Bold', fontSize=12,
            spaceBefore=14, spaceAfter=6, textColor=HexColor('#2a2f55'),
        ),
        'Meta': ParagraphStyle(
            'DiscussMeta', parent=base['BodyText'],
            fontName='Helvetica', fontSize=9.5,
            leading=13, textColor=HexColor('#444'),
        ),
        'Body': ParagraphStyle(
            'DiscussBody', parent=base['BodyText'],
            fontName='Helvetica', fontSize=10.5, leading=15,
            spaceAfter=6,
        ),
        'Error': ParagraphStyle(
            'DiscussError', parent=base['BodyText'],
            fontName='Helvetica-Oblique', fontSize=10,
            textColor=HexColor('#b00020'), leading=14,
        ),
        'Moderator': ParagraphStyle(
            'DiscussModerator', parent=base['BodyText'],
            fontName='Helvetica-Oblique', fontSize=10,
            textColor=HexColor('#555588'), leading=13,
            spaceBefore=4, spaceAfter=2,
        ),
    }
    return styles


def _escape_pdf(text):
    """Reportlab Paragraph parses an XML-like subset; escape the triggers."""
    return (
        (text or '')
        .replace('&', '&amp;')
        .replace('<', '&lt;')
        .replace('>', '&gt;')
    )


def session_to_pdf(session):
    """Render the session to a PDF file and return the bytes."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        leftMargin=0.75 * inch, rightMargin=0.75 * inch,
        topMargin=0.75 * inch, bottomMargin=0.75 * inch,
        title=session.title or 'Multi-Agent Discussion',
    )
    styles = _pdf_styles()
    story = []

    # Title
    story.append(Paragraph(_escape_pdf(session.title or 'Multi-Agent Discussion'), styles['Title']))
    story.append(HRFlowable(width='100%', thickness=0.6, color=HexColor('#c8cce0'), spaceBefore=0, spaceAfter=8))

    # Metadata
    for label, value in _metadata_lines(session):
        story.append(Paragraph(
            f'<b>{_escape_pdf(label)}:</b> {_escape_pdf(str(value))}',
            styles['Meta'],
        ))
    story.append(Spacer(1, 10))

    # Problem
    story.append(Paragraph('Problem Statement', styles['H1']))
    problem_text = _escape_pdf(session.initial_problem or '').replace('\n', '<br/>')
    story.append(Paragraph(problem_text, styles['Body']))

    # Discussion
    story.append(Paragraph('Discussion', styles['H1']))
    turns = _iter_turns_in_order(session)
    if not turns:
        story.append(Paragraph('(No turns recorded.)', styles['Body']))
    else:
        for turn in turns:
            name_lower = (turn.model_name or '').lower()
            is_user = name_lower == 'user'
            is_moderator = name_lower == 'moderator'

            if is_moderator:
                chosen = (turn.model_role or '').lstrip('->').strip()
                header = f'Moderator -> {chosen}' if chosen else 'Moderator'
                story.append(Paragraph(
                    f'<b>{_escape_pdf(header)}</b>',
                    styles['Moderator'],
                ))
                if turn.content and turn.content.strip():
                    story.append(Paragraph(
                        _escape_pdf(turn.content.strip()).replace('\n', '<br/>'),
                        styles['Moderator'],
                    ))
                continue

            label = 'Human (user)' if is_user else (turn.model_name or 'unknown')
            heading_bits = [f'Turn {turn.turn_number}', label]
            if turn.model_role and not is_user:
                heading_bits.append(f'role: {turn.model_role}')
            if turn.duration is not None and not is_user:
                heading_bits.append(f'{turn.duration:.2f}s')
            story.append(Paragraph(
                _escape_pdf(' · '.join(heading_bits)),
                styles['Turn'],
            ))
            if turn.error:
                story.append(Paragraph(
                    f'[error] {_escape_pdf(turn.error)}',
                    styles['Error'],
                ))
            else:
                text = turn.content or ''
                for block in text.split('\n\n'):
                    block = block.strip()
                    if not block:
                        continue
                    story.append(Paragraph(
                        _escape_pdf(block).replace('\n', '<br/>'),
                        styles['Body'],
                    ))

    # Synthesis section (if available)
    if session.synthesis:
        story.append(Paragraph('Synthesis', styles['H1']))
        if session.synthesis_model:
            story.append(Paragraph(
                f'<i>Synthesized by: {_escape_pdf(session.synthesis_model)}</i>',
                styles['Meta'],
            ))
        for block in session.synthesis.split('\n\n'):
            block = block.strip()
            if not block:
                continue
            story.append(Paragraph(
                _escape_pdf(block).replace('\n', '<br/>'),
                styles['Body'],
            ))

    doc.build(story)
    return buf.getvalue()
